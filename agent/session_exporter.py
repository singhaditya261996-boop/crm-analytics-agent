"""
agent/session_exporter.py — Comprehensive export system for CRM Analytics Agent (Module 10).

Six export types:
  1. Full session Word doc  (export_session_docx)
  2. Per-section Word doc   (export_section_docx)
  3. [Excel pivot — handled by PivotHandler]
  4. Chat history text      (export_chat_txt)
  5. Actions list Word doc  (export_actions_docx)
  6. Dashboard PDF          (export_dashboard_pdf)

All output is written to exports/ (or caller-supplied exports_dir).
Zero network calls — 100% local.
"""
from __future__ import annotations

import asyncio
import io
import logging
import tempfile
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd

logger = logging.getLogger(__name__)

# ── Styling constants ─────────────────────────────────────────────────────────

_NAVY        = "1A2A5E"   # H1, table headers
_TEAL        = "0D9488"   # H2
_TEAL_LIGHT  = "CCFBF1"   # recommendation box fill
_ALT_ROW     = "E8F0FE"   # alternating table rows
_WHITE       = "FFFFFF"
_CONFIDENTIALITY = "Prepared under NDA — not for external distribution"

# ── Module-level python-docx helpers ─────────────────────────────────────────


def _pt(points: int) -> Any:
    """Return a docx Pt length object."""
    from docx.shared import Pt
    return Pt(points)


def _rgb(hex_str: str) -> Any:
    """Return docx RGBColor from a 6-char hex string."""
    from docx.shared import RGBColor
    r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def _cell_fill(cell: Any, hex_str: str) -> None:
    """Set background colour on a python-docx table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str)
    tc_pr.append(shd)


def _h1(doc: Any, text: str) -> Any:
    """Add a navy H1 heading."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = _pt(18)
    run.font.color.rgb = _rgb(_NAVY)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _h2(doc: Any, text: str) -> Any:
    """Add a teal H2 heading."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = _pt(14)
    run.font.color.rgb = _rgb(_TEAL)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_rec_box(doc: Any, rec: dict) -> None:
    """Add a teal recommendation box with priority action, risk flag, opportunity."""
    if not rec or not any(rec.values()):
        return
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _cell_fill(cell, _TEAL_LIGHT)
    lines = []
    if rec.get("priority_action"):
        lines.append(f"\U0001f3af Priority Action: {rec['priority_action']}")
    if rec.get("risk_flag"):
        lines.append(f"\u26a0\ufe0f Risk Flag: {rec['risk_flag']}")
    if rec.get("opportunity"):
        lines.append(f"\U0001f4a1 Opportunity: {rec['opportunity']}")
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = para.add_run(line)
        run.bold = (i == 0)
        run.font.color.rgb = _rgb(_TEAL)
    doc.add_paragraph()


def _add_conf_footer(doc: Any, score: int) -> None:
    """Add a confidence note paragraph (green ≥85 / amber ≥60 / red <60)."""
    if score >= 85:
        label, colour = "High confidence", "16A34A"
    elif score >= 60:
        label, colour = "Review recommended", "CA8A04"
    else:
        label, colour = "Low confidence — verify manually", "DC2626"
    p = doc.add_paragraph()
    run = p.add_run(f"Confidence: {label} ({score}/100)")
    run.italic = True
    run.font.size = _pt(9)
    run.font.color.rgb = _rgb(colour)


def _add_data_table(doc: Any, df: pd.DataFrame, max_rows: int = 10) -> None:
    """Add a formatted data table — navy header, alternating rows."""
    if df is None or df.empty:
        return
    display = df.head(max_rows).reset_index(drop=True)
    cols = list(display.columns)
    tbl = doc.add_table(rows=1 + len(display), cols=len(cols))
    tbl.style = "Table Grid"
    # Header
    hdr = tbl.rows[0]
    for c_idx, col_name in enumerate(cols):
        cell = hdr.cells[c_idx]
        _cell_fill(cell, _NAVY)
        run = cell.paragraphs[0].add_run(str(col_name))
        run.bold = True
        run.font.color.rgb = _rgb(_WHITE)
        run.font.size = _pt(9)
    # Data rows
    for r_idx, row_data in enumerate(display.itertuples(index=False), start=1):
        fill = _ALT_ROW if r_idx % 2 == 0 else None
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx].cells[c_idx]
            if fill:
                _cell_fill(cell, fill)
            cell.paragraphs[0].add_run(str(val) if pd.notna(val) else "—").font.size = _pt(9)
    doc.add_paragraph()


def _embed_chart(doc: Any, fig: Any) -> None:
    """Save a plotly figure as a PNG (via kaleido) and embed it centred."""
    try:
        import kaleido  # noqa: F401
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = Path(tmp.name)
        fig.write_image(str(tmp_path), width=1200, height=600, scale=2)
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(str(tmp_path), width=Inches(6))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        try:
            tmp_path.unlink()
        except OSError:
            pass
    except Exception as exc:
        logger.warning("Chart embedding skipped: %s", exc)


def _add_confidentiality_footer(section: Any) -> None:
    """Add a light-grey confidentiality line to a section/paragraph."""
    p = section.add_paragraph(_CONFIDENTIALITY)
    run = p.runs[0]
    run.italic = True
    run.font.size = _pt(8)
    run.font.color.rgb = _rgb("94A3B8")


# ── SessionExporter ───────────────────────────────────────────────────────────


class SessionExporter:
    """
    Comprehensive export system for the CRM Analytics Agent.

    Parameters
    ----------
    exports_dir : Path, optional
        Root output directory. Defaults to ``exports/``.
    llm_client : Any, optional
        LLMClient instance used to generate executive summaries.
        If None, a template-based summary is used instead.
    """

    def __init__(
        self,
        exports_dir: Path | None = None,
        llm_client: Any = None,
    ) -> None:
        self.exports_dir = Path(exports_dir) if exports_dir else Path("exports")
        self.exports_dir.mkdir(parents=True, exist_ok=True)
        self._llm = llm_client

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _ts(self) -> str:
        return datetime.now().strftime("%Y%m%d_%H%M%S")

    def _run_async(self, coro: Any) -> Any:
        """Run an async coroutine from a sync context."""
        try:
            loop = asyncio.get_running_loop()
        except RuntimeError:
            loop = None
        if loop and loop.is_running():
            with ThreadPoolExecutor(max_workers=1) as pool:
                return pool.submit(asyncio.run, coro).result()
        return asyncio.run(coro)

    def _executive_summary(
        self,
        agenda_results: dict,
        session_id: str = "",
        model_used: str = "",
    ) -> str:
        """
        Generate a 200-word executive summary from agenda results.
        Uses LLM if available, else produces a template-based fallback.
        """
        # Collect top priority actions
        actions: list[str] = []
        risks: list[str] = []
        opps: list[str] = []
        n_items = 0
        sections_covered: set[int] = set()
        for sec_num, sec_results in agenda_results.items():
            sections_covered.add(sec_num)
            for _q, result in sec_results.items():
                n_items += 1
                rec = getattr(result, "recommendation", {}) or {}
                if rec.get("priority_action"):
                    actions.append(rec["priority_action"])
                if rec.get("risk_flag"):
                    risks.append(rec["risk_flag"])
                if rec.get("opportunity"):
                    opps.append(rec["opportunity"])

        # Template fallback (used when llm_client is None or LLM call fails)
        def _template() -> str:
            act_text = "; ".join(actions[:3]) if actions else "see individual sections"
            risk_text = "; ".join(risks[:3]) if risks else "none flagged"
            lines = [
                f"This session covered {n_items} agenda items across "
                f"{len(sections_covered)} sections.",
                f"Key findings: {act_text}.",
                f"Top risks: {risk_text}.",
            ]
            if opps:
                lines.append(f"Opportunities identified: {'; '.join(opps[:2])}.")
            if model_used:
                lines.append(f"Analysis performed locally using model: {model_used}.")
            if session_id:
                lines.append(f"Session ID: {session_id}.")
            return " ".join(lines)

        if self._llm is None:
            return _template()

        # LLM-based summary
        top_actions = "\n".join(f"- {a}" for a in actions[:5])
        top_risks   = "\n".join(f"- {r}" for r in risks[:3])
        prompt = (
            f"Write a 200-word executive summary for a CRM analytics session report.\n"
            f"Items covered: {n_items} across {len(sections_covered)} sections.\n"
            f"Priority actions:\n{top_actions}\n"
            f"Key risks:\n{top_risks}\n"
            "Use professional business English. Be concise and direct."
        )
        try:
            return self._run_async(
                self._llm.complete("You are an executive report writer.", prompt)
            )
        except Exception as exc:
            logger.warning("LLM executive summary failed (%s) — using template", exc)
            return _template()

    def _build_dashboard_figs(
        self, dataframes: dict[str, pd.DataFrame]
    ) -> list[Any]:
        """
        Build the same 6 plotly figures as the dashboard tab.
        Returns an empty list if plotly is not installed.
        """
        try:
            import plotly.graph_objects as go
        except ImportError:
            return []

        dfs = dataframes
        if not dfs:
            return []

        figs: list[Any] = []

        def _find(df: pd.DataFrame, keywords: list[str]) -> str | None:
            for kw in keywords:
                for col in df.columns:
                    if kw in col.lower():
                        return col
            return None

        def _pick(table_kws: list[str]) -> tuple[str, pd.DataFrame]:
            for kw in table_kws:
                for name, df in dfs.items():
                    if kw in name.lower():
                        return name, df
            return next(iter(dfs.items()))

        # 1. Revenue Pareto
        try:
            _, df_a = _pick(["account", "client", "customer"])
            rev = _find(df_a, ["revenue", "amount", "value", "deal_value"])
            nam = _find(df_a, ["account_name", "name", "client_name", "company"])
            if rev and nam:
                top = (
                    df_a.groupby(nam)[rev].sum()
                    .sort_values(ascending=False).head(20).reset_index()
                )
                total = top[rev].sum()
                top["_cum"] = top[rev].cumsum() / total * 100 if total > 0 else 0
                fig = go.Figure()
                fig.add_trace(go.Bar(x=top[nam], y=top[rev], name="Revenue"))
                fig.add_trace(go.Scatter(
                    x=top[nam], y=top["_cum"], name="Cum %",
                    yaxis="y2", mode="lines+markers",
                    line=dict(color="orange"),
                ))
                fig.update_layout(
                    title="Revenue Concentration (Pareto)",
                    yaxis2=dict(overlaying="y", side="right", range=[0, 110]),
                    height=500, width=1100,
                )
                figs.append(fig)
        except Exception as exc:
            logger.debug("Revenue pareto chart skipped: %s", exc)

        # 2. Pipeline Funnel
        try:
            _, df_b = _pick(["opportunit", "pipeline", "deal"])
            stage = _find(df_b, ["stage", "status", "phase"])
            val   = _find(df_b, ["value", "amount", "deal_value", "revenue"])
            if stage:
                if val:
                    fd = df_b.groupby(stage)[val].sum().reset_index()
                    fd = fd.sort_values(val, ascending=False)
                    y_vals, x_vals = fd[stage].astype(str).tolist(), fd[val].tolist()
                else:
                    vc = df_b[stage].value_counts()
                    y_vals, x_vals = vc.index.astype(str).tolist(), vc.values.tolist()
                fig = go.Figure(go.Funnel(
                    y=y_vals, x=x_vals, textinfo="value+percent initial"
                ))
                fig.update_layout(title="Pipeline Stage Funnel", height=500, width=1100)
                figs.append(fig)
        except Exception as exc:
            logger.debug("Pipeline funnel chart skipped: %s", exc)

        # 3. Win Rate by Service Line
        try:
            _, df_c = _pick(["opportunit", "pipeline", "deal"])
            svc = _find(df_c, ["service_line", "service", "division", "category"])
            stg = _find(df_c, ["stage", "status", "outcome"])
            if svc and stg:
                won = df_c[df_c[stg].str.lower().isin(["won", "closed won", "win"])]
                total_s = df_c.groupby(svc).size()
                won_s   = won.groupby(svc).size()
                wr = (won_s / total_s * 100).fillna(0).sort_values(
                    ascending=False).reset_index()
                wr.columns = [svc, "win_rate_pct"]
                fig = go.Figure(go.Bar(
                    x=wr[svc].astype(str), y=wr["win_rate_pct"],
                    text=(wr["win_rate_pct"].round(1).astype(str) + "%"),
                    textposition="outside",
                ))
                fig.add_hline(y=28, line_dash="dash", line_color="orange",
                              annotation_text="FM benchmark 28%")
                fig.update_layout(
                    title="Win Rate by Service Line",
                    yaxis_title="Win Rate %", height=500, width=1100,
                    yaxis_range=[0, min(100, wr["win_rate_pct"].max() * 1.25 + 5)],
                )
                figs.append(fig)
        except Exception as exc:
            logger.debug("Win rate chart skipped: %s", exc)

        # 4. Sales Cycle Trend
        try:
            _, df_d = _pick(["opportunit", "pipeline", "deal"])
            close_c = _find(df_d, ["close_date", "closed_date", "close", "end_date"])
            open_c  = _find(df_d, ["created_date", "created_at", "start_date", "open_date"])
            if close_c and open_c:
                dc = df_d[[close_c, open_c]].dropna().copy()
                dc[close_c] = pd.to_datetime(dc[close_c], errors="coerce")
                dc[open_c]  = pd.to_datetime(dc[open_c],  errors="coerce")
                dc["days"]  = (dc[close_c] - dc[open_c]).dt.days
                dc = dc.dropna(subset=["days"])
                dc["month"] = dc[close_c].dt.to_period("M").astype(str)
                monthly = dc.groupby("month")["days"].mean().reset_index()
                monthly.columns = ["month", "avg_days"]
                monthly["_ma"] = monthly["avg_days"].rolling(3, min_periods=1).mean()
                fig = go.Figure()
                fig.add_trace(go.Scatter(
                    x=monthly["month"], y=monthly["avg_days"],
                    mode="lines+markers", name="Avg days", opacity=0.5,
                ))
                fig.add_trace(go.Scatter(
                    x=monthly["month"], y=monthly["_ma"],
                    mode="lines", name="3-month MA", line=dict(width=2),
                ))
                fig.add_hline(y=84, line_dash="dash", line_color="orange",
                              annotation_text="FM benchmark 84d")
                fig.update_layout(
                    title="Sales Cycle Trend",
                    yaxis_title="Days", height=500, width=1100,
                )
                figs.append(fig)
        except Exception as exc:
            logger.debug("Sales cycle chart skipped: %s", exc)

        # 5. Top 10 Accounts by Revenue
        try:
            _, df_t = _pick(["account", "client", "customer"])
            rev = _find(df_t, ["revenue", "amount", "value", "deal_value"])
            nam = _find(df_t, ["account_name", "name", "client_name", "company"])
            if rev and nam:
                top10 = (
                    df_t.groupby(nam)[rev].sum()
                    .sort_values(ascending=True).tail(10).reset_index()
                )
                fig = go.Figure(go.Bar(
                    x=top10[rev], y=top10[nam].astype(str),
                    orientation="h",
                    text=top10[rev].apply(
                        lambda v: f"£{v:,.0f}" if pd.notna(v) else "—"
                    ),
                    textposition="outside",
                ))
                fig.update_layout(
                    title="Top 10 Accounts by Revenue",
                    xaxis_title="Revenue", height=500, width=1100,
                )
                figs.append(fig)
        except Exception as exc:
            logger.debug("Top 10 accounts chart skipped: %s", exc)

        # 6. Account Health Heatmap
        try:
            _, df_h = _pick(["account", "client", "customer"])
            num_cols = df_h.select_dtypes(include="number").columns.tolist()
            if len(num_cols) >= 2:
                heat = df_h[num_cols].head(30).fillna(0)
                nam  = _find(df_h, ["account_name", "name", "client_name"])
                y_labels = (
                    df_h[nam].head(30).astype(str).tolist()
                    if nam else [str(i) for i in range(len(heat))]
                )
                fig = go.Figure(go.Heatmap(
                    z=heat.values, x=heat.columns.tolist(),
                    y=y_labels, colorscale="Blues",
                ))
                fig.update_layout(
                    title="Account Health Heatmap", height=600, width=1100,
                )
                figs.append(fig)
        except Exception as exc:
            logger.debug("Heatmap chart skipped: %s", exc)

        return figs

    # ── Export type 1 — Full session Word doc ─────────────────────────────────

    def export_session_docx(
        self,
        chat_history: list[dict],
        agenda_results: dict[int, dict[str, Any]],
        dataframes: dict[str, pd.DataFrame],
        profiles: dict,
        session_id: str = "",
        session_start: datetime | None = None,
        model_used: str = "",
    ) -> Path:
        """
        Export a full session to a professionally formatted Word document.

        Structure:
          Cover page → Executive Summary → Per-section results →
          Actions & Next Steps → Data Quality Notes → Appendix
        """
        try:
            import docx
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise ImportError("pip install python-docx") from exc

        doc = docx.Document()

        # ── Cover page ────────────────────────────────────────────────────────
        cover = doc.add_paragraph()
        cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cover.add_run("CRM Analytics — Weekly Review")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = _rgb(_NAVY)

        date_str = (session_start or datetime.now()).strftime("%d %B %Y")
        doc.add_paragraph(f"Date: {date_str}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        if session_id:
            doc.add_paragraph(f"Session ID: {session_id}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_p = doc.add_paragraph(_CONFIDENTIALITY)
        conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_p.runs[0].italic = True
        conf_p.runs[0].font.color.rgb = _rgb("94A3B8")
        doc.add_page_break()

        # ── Executive Summary ─────────────────────────────────────────────────
        _h1(doc, "Executive Summary")
        summary_text = self._executive_summary(agenda_results, session_id, model_used)
        doc.add_paragraph(summary_text)

        # Top Priority Actions
        actions = [
            result.recommendation.get("priority_action", "")
            for sec_results in agenda_results.values()
            for result in sec_results.values()
            if getattr(result, "recommendation", {}).get("priority_action")
        ]
        if actions:
            doc.add_paragraph()
            _h2(doc, "Top Priority Actions")
            for i, act in enumerate(actions[:6], 1):
                doc.add_paragraph(f"{i}. {act}", style="List Number")

        # Top Risks / Opportunities
        risks = [
            result.recommendation.get("risk_flag", "")
            for sec_results in agenda_results.values()
            for result in sec_results.values()
            if getattr(result, "recommendation", {}).get("risk_flag")
        ]
        opps = [
            result.recommendation.get("opportunity", "")
            for sec_results in agenda_results.values()
            for result in sec_results.values()
            if getattr(result, "recommendation", {}).get("opportunity")
        ]
        if risks or opps:
            doc.add_paragraph()
            _h2(doc, "Risks & Opportunities")
            for r in risks[:3]:
                doc.add_paragraph(f"\u26a0\ufe0f {r}", style="List Bullet")
            for o in opps[:3]:
                doc.add_paragraph(f"\U0001f4a1 {o}", style="List Bullet")
        doc.add_page_break()

        # ── Per-agenda-section results ─────────────────────────────────────────
        from agenda.prompts import SECTION_TITLES
        for sec_num in sorted(agenda_results.keys()):
            sec_results = agenda_results[sec_num]
            title = SECTION_TITLES.get(sec_num, f"Section {sec_num}")
            _h1(doc, f"Section {sec_num} — {title}")

            for question, result in sec_results.items():
                _h2(doc, question)

                if getattr(result, "answer_text", ""):
                    doc.add_paragraph(result.answer_text)

                rec = getattr(result, "recommendation", {}) or {}
                _add_rec_box(doc, rec)

                if getattr(result, "benchmark_used", False):
                    bm_p = doc.add_paragraph()
                    run = bm_p.add_run("\U0001f4ca Benchmark data was used in this analysis.")
                    run.italic = True
                    run.font.color.rgb = _rgb(_TEAL)

                if getattr(result, "chart", None) is not None:
                    _embed_chart(doc, result.chart)

                if getattr(result, "result_df", None) is not None:
                    _add_data_table(doc, result.result_df)
                elif getattr(result, "pivot_df", None) is not None:
                    _add_data_table(doc, result.pivot_df)

                _add_conf_footer(doc, getattr(result, "confidence_score", 0))
                doc.add_paragraph("─" * 60)

            doc.add_page_break()

        # ── Chat history section (if no agenda results) ───────────────────────
        if chat_history and not agenda_results:
            _h1(doc, "Chat Session")
            for msg in chat_history:
                if msg["role"] == "user":
                    p = doc.add_paragraph()
                    p.add_run(f"Q: {msg['content']}").bold = True
                elif msg.get("result"):
                    r = msg["result"]
                    doc.add_paragraph(f"A: {getattr(r, 'answer_text', '')}")
                    _add_rec_box(doc, getattr(r, "recommendation", {}) or {})
                    _add_conf_footer(doc, getattr(r, "confidence_score", 0))
            doc.add_page_break()

        # ── Actions & Next Steps ───────────────────────────────────────────────
        _h1(doc, "Actions & Next Steps")
        action_rows: list[dict] = []
        for sec_num, sec_results in agenda_results.items():
            for question, result in sec_results.items():
                rec = getattr(result, "recommendation", {}) or {}
                action_rows.append({
                    "Section": sec_num,
                    "Priority Action": rec.get("priority_action", "—"),
                    "Source Question": question[:80],
                    "Risk Flag": rec.get("risk_flag", "—"),
                    "Opportunity": rec.get("opportunity", "—"),
                    "Owner": "",
                    "Due Date": "",
                })

        if action_rows:
            cols = ["Section", "Priority Action", "Source Question",
                    "Risk Flag", "Opportunity", "Owner", "Due Date"]
            tbl = doc.add_table(rows=1 + len(action_rows), cols=len(cols))
            tbl.style = "Table Grid"
            hdr = tbl.rows[0]
            for c_idx, col_name in enumerate(cols):
                cell = hdr.cells[c_idx]
                _cell_fill(cell, _NAVY)
                run = cell.paragraphs[0].add_run(col_name)
                run.bold = True
                run.font.color.rgb = _rgb(_WHITE)
                run.font.size = Pt(8)
            for r_idx, row_data in enumerate(action_rows, 1):
                fill = _ALT_ROW if r_idx % 2 == 0 else None
                for c_idx, col_name in enumerate(cols):
                    cell = tbl.rows[r_idx].cells[c_idx]
                    if fill:
                        _cell_fill(cell, fill)
                    cell.paragraphs[0].add_run(str(row_data[col_name])).font.size = Pt(8)
        doc.add_page_break()

        # ── Data Quality Notes ─────────────────────────────────────────────────
        dq_results = agenda_results.get(5, {})
        if dq_results:
            _h1(doc, "Data Quality Notes")
            for question, result in dq_results.items():
                _h2(doc, question)
                if getattr(result, "answer_text", ""):
                    doc.add_paragraph(result.answer_text)
            doc.add_page_break()

        # ── Appendix ──────────────────────────────────────────────────────────
        _h1(doc, "Appendix")

        # All pandas queries
        _h2(doc, "A. Generated Pandas Queries")
        for sec_num, sec_results in agenda_results.items():
            for question, result in sec_results.items():
                code = getattr(result, "code", "")
                if code:
                    p = doc.add_paragraph()
                    p.add_run(f"Q: {question[:80]}").bold = True
                    doc.add_paragraph(code, style="No Spacing")
                    doc.add_paragraph()

        # Data sources
        _h2(doc, "B. Data Sources")
        for tbl_name, df in dataframes.items():
            doc.add_paragraph(f"• {tbl_name}: {len(df):,} rows × {len(df.columns)} columns")

        # Session metadata
        _h2(doc, "C. Session Metadata")
        all_results = [
            r for sec in agenda_results.values() for r in sec.values()
        ]
        n_queries = len(all_results)
        avg_conf = (
            sum(getattr(r, "confidence_score", 0) for r in all_results) / n_queries
            if n_queries else 0
        )
        doc.add_paragraph(f"Session ID: {session_id}")
        doc.add_paragraph(f"Model used: {model_used}")
        doc.add_paragraph(f"Total queries: {n_queries}")
        doc.add_paragraph(f"Average confidence score: {avg_conf:.1f}/100")
        doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).isoformat()}")
        _add_confidentiality_footer(doc)

        # Save
        out_dir = self.exports_dir / "sessions"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"session_{self._ts()}.docx"
        doc.save(str(out_path))
        logger.info("Full session exported to %s", out_path)
        return out_path

    # ── Export type 2 — Per-section Word doc ─────────────────────────────────

    def export_section_docx(
        self,
        section_num: int,
        results: dict[str, Any],
    ) -> Path:
        """
        Export a single agenda section to a Word document.
        """
        try:
            import docx
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise ImportError("pip install python-docx") from exc

        from agenda.prompts import SECTION_TITLES
        title = SECTION_TITLES.get(section_num, f"Section {section_num}")

        doc = docx.Document()

        # Cover
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Equans CRM Analytics — {title}")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = _rgb(_NAVY)
        doc.add_paragraph(
            datetime.now().strftime("%d %B %Y")
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for question, result in results.items():
            _h2(doc, question)
            if getattr(result, "answer_text", ""):
                doc.add_paragraph(result.answer_text)
            rec = getattr(result, "recommendation", {}) or {}
            _add_rec_box(doc, rec)
            if getattr(result, "chart", None) is not None:
                _embed_chart(doc, result.chart)
            if getattr(result, "result_df", None) is not None:
                _add_data_table(doc, result.result_df)
            elif getattr(result, "pivot_df", None) is not None:
                _add_data_table(doc, result.pivot_df)
            _add_conf_footer(doc, getattr(result, "confidence_score", 0))
            doc.add_paragraph()

        _add_confidentiality_footer(doc)

        out_dir = self.exports_dir / "sessions"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"section_{section_num}_{self._ts()}.docx"
        doc.save(str(out_path))
        logger.info("Section %d exported to %s", section_num, out_path)
        return out_path

    # ── Export type 4 — Chat history text ────────────────────────────────────

    def export_chat_txt(
        self,
        chat_history: list[dict],
        session_id: str = "",
    ) -> Path:
        """
        Export chat history to a plain-text file.

        Format::

            === CRM Analytics Agent — Chat Export ===
            Date: YYYY-MM-DD
            Session ID: <id>

            [HH:MM] USER: <question>
            [HH:MM] AGENT: <answer>
                     RECOMMENDATION: Priority Action: ...
        """
        lines: list[str] = [
            "=== CRM Analytics Agent — Chat Export ===",
            f"Date: {datetime.now().strftime('%Y-%m-%d')}",
        ]
        if session_id:
            lines.append(f"Session ID: {session_id}")
        lines.append(_CONFIDENTIALITY)
        lines.append("")

        for msg in chat_history:
            role = msg.get("role", "")
            content = msg.get("content", "")
            ts = msg.get("timestamp", "")
            ts_str = ""
            if ts:
                if isinstance(ts, datetime):
                    ts_str = f"[{ts.strftime('%H:%M')}] "
                else:
                    ts_str = f"[{str(ts)[:5]}] "

            if role == "user":
                lines.append(f"{ts_str}USER: {content}")
            elif role == "assistant":
                lines.append(f"{ts_str}AGENT: {content}")
                result = msg.get("result")
                if result is not None:
                    rec = getattr(result, "recommendation", {}) or {}
                    if rec.get("priority_action"):
                        lines.append(f"         RECOMMENDATION: Priority Action: {rec['priority_action']}")
                    if rec.get("risk_flag"):
                        lines.append(f"         Risk Flag: {rec['risk_flag']}")
                    if rec.get("opportunity"):
                        lines.append(f"         Opportunity: {rec['opportunity']}")
                    score = getattr(result, "confidence_score", None)
                    if score is not None:
                        if score >= 85:
                            label = "\U0001f7e2 High"
                        elif score >= 60:
                            label = "\U0001f7e1 Review"
                        else:
                            label = "\U0001f534 Low"
                        lines.append(f"         Confidence: {label} ({score}/100)")
                lines.append("")

        out_dir = self.exports_dir / "sessions"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"chat_{self._ts()}.txt"
        out_path.write_text("\n".join(lines), encoding="utf-8")
        logger.info("Chat history exported to %s", out_path)
        return out_path

    # ── Export type 5 — Actions list Word doc ─────────────────────────────────

    def export_actions_docx(
        self,
        agenda_results: dict[int, dict[str, Any]],
    ) -> Path:
        """
        Export a consolidated actions & next steps document.

        Table columns: # · Priority Action · Source Question ·
                        Risk Flag · Opportunity · Owner · Due Date
        """
        try:
            import docx
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise ImportError("pip install python-docx") from exc

        from agenda.prompts import SECTION_TITLES

        doc = docx.Document()

        # Cover
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Equans CRM — Actions & Priorities")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = _rgb(_NAVY)
        doc.add_paragraph(
            datetime.now().strftime("%d %B %Y")
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Build rows ordered by section
        rows: list[dict] = []
        for sec_num in sorted(agenda_results.keys()):
            sec_results = agenda_results[sec_num]
            for question, result in sec_results.items():
                rec = getattr(result, "recommendation", {}) or {}
                rows.append({
                    "#": len(rows) + 1,
                    "Priority Action": rec.get("priority_action", "—"),
                    "Source Question": question[:80],
                    "Risk Flag": rec.get("risk_flag", "—"),
                    "Opportunity": rec.get("opportunity", "—"),
                    "Owner": "",
                    "Due Date": "",
                    "_section": sec_num,
                })

        if not rows:
            doc.add_paragraph("No agenda results found for this session.")
            _add_confidentiality_footer(doc)
            out_dir = self.exports_dir / "sessions"
            out_dir.mkdir(parents=True, exist_ok=True)
            out_path = out_dir / f"actions_{self._ts()}.docx"
            doc.save(str(out_path))
            return out_path

        cols = ["#", "Priority Action", "Source Question",
                "Risk Flag", "Opportunity", "Owner", "Due Date"]

        # If >10 rows: group by section with H2 subheadings
        if len(rows) > 10:
            current_section: int | None = None
            tbl: Any = None
            row_num = 0
            for row_data in rows:
                sec_n = row_data["_section"]
                if sec_n != current_section:
                    if tbl is not None:
                        doc.add_paragraph()
                    sec_title = SECTION_TITLES.get(sec_n, f"Section {sec_n}")
                    _h2(doc, f"Section {sec_n} — {sec_title}")
                    tbl = doc.add_table(rows=1, cols=len(cols))
                    tbl.style = "Table Grid"
                    hdr = tbl.rows[0]
                    for c_idx, col_name in enumerate(cols):
                        cell = hdr.cells[c_idx]
                        _cell_fill(cell, _NAVY)
                        run = cell.paragraphs[0].add_run(col_name)
                        run.bold = True
                        run.font.color.rgb = _rgb(_WHITE)
                        run.font.size = Pt(8)
                    current_section = sec_n
                    row_num = 0
                row_num += 1
                new_row = tbl.add_row()
                fill = _ALT_ROW if row_num % 2 == 0 else None
                for c_idx, col_name in enumerate(cols):
                    cell = new_row.cells[c_idx]
                    if fill:
                        _cell_fill(cell, fill)
                    cell.paragraphs[0].add_run(str(row_data[col_name])).font.size = Pt(8)
        else:
            tbl = doc.add_table(rows=1 + len(rows), cols=len(cols))
            tbl.style = "Table Grid"
            hdr = tbl.rows[0]
            for c_idx, col_name in enumerate(cols):
                cell = hdr.cells[c_idx]
                _cell_fill(cell, _NAVY)
                run = cell.paragraphs[0].add_run(col_name)
                run.bold = True
                run.font.color.rgb = _rgb(_WHITE)
                run.font.size = Pt(8)
            for r_idx, row_data in enumerate(rows, 1):
                fill = _ALT_ROW if r_idx % 2 == 0 else None
                for c_idx, col_name in enumerate(cols):
                    cell = tbl.rows[r_idx].cells[c_idx]
                    if fill:
                        _cell_fill(cell, fill)
                    cell.paragraphs[0].add_run(str(row_data[col_name])).font.size = Pt(8)

        _add_confidentiality_footer(doc)

        out_dir = self.exports_dir / "sessions"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"actions_{self._ts()}.docx"
        doc.save(str(out_path))
        logger.info("Actions doc exported to %s", out_path)
        return out_path

    # ── Export type 6 — Dashboard PDF ────────────────────────────────────────

    def export_dashboard_pdf(
        self,
        dataframes: dict[str, pd.DataFrame],
    ) -> Path:
        """
        Export the dashboard charts as a multi-page PDF.

        Uses kaleido to render each plotly figure as a PDF page, then
        merges all pages with pypdf.
        """
        try:
            import kaleido  # noqa: F401
        except ImportError as exc:
            raise ImportError("pip install kaleido") from exc
        try:
            from pypdf import PdfWriter
        except ImportError as exc:
            raise ImportError("pip install pypdf") from exc

        figs = self._build_dashboard_figs(dataframes)
        if not figs:
            raise ValueError("No charts could be generated — check dataframe contents.")

        tmp_paths: list[Path] = []
        try:
            for i, fig in enumerate(figs):
                with tempfile.NamedTemporaryFile(
                    suffix=".pdf", delete=False, prefix=f"dash_{i}_"
                ) as tmp:
                    tmp_path = Path(tmp.name)
                fig.write_image(str(tmp_path), format="pdf", width=1200, height=700)
                tmp_paths.append(tmp_path)

            writer = PdfWriter()
            for tmp_path in tmp_paths:
                writer.append(str(tmp_path))

            out_dir = self.exports_dir
            out_dir.mkdir(parents=True, exist_ok=True)
            out_path = out_dir / f"dashboard_{datetime.now().strftime('%Y%m%d')}.pdf"
            with open(out_path, "wb") as fh:
                writer.write(fh)
        finally:
            for tmp_path in tmp_paths:
                try:
                    tmp_path.unlink()
                except OSError:
                    pass

        logger.info("Dashboard PDF exported to %s", out_path)
        return out_path
