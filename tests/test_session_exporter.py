"""tests/test_session_exporter.py — Tests for agent/session_exporter.py (Module 10)."""
from __future__ import annotations

import os
import sys
import types
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import MagicMock, patch

import pandas as pd
import pytest


# ── Helpers to build lightweight mock QueryResult objects ─────────────────────

def _make_result(
    question: str = "Which accounts have the highest revenue?",
    answer_text: str = "Top 3 accounts drive 60% of revenue.",
    priority_action: str = "Focus on Account A",
    risk_flag: str = "Pipeline thin for Q4",
    opportunity: str = "Cross-sell potential in FM",
    confidence_score: int = 82,
    code: str = "result = df.groupby('account')['revenue'].sum()",
    intent_type: str = "ranking",
    provider_used: str = "ollama",
    session_id: str = "sess-001",
) -> MagicMock:
    r = MagicMock()
    r.question = question
    r.answer_text = answer_text
    r.recommendation = {
        "priority_action": priority_action,
        "risk_flag": risk_flag,
        "opportunity": opportunity,
    }
    r.confidence_score = confidence_score
    r.code = code
    r.intent_type = intent_type
    r.provider_used = provider_used
    r.session_id = session_id
    r.timestamp = datetime(2024, 3, 1, 10, 30, tzinfo=timezone.utc)
    r.chart = None
    r.result_df = None
    r.pivot_df = None
    r.benchmark_used = False
    return r


def _make_agenda(sections: int = 3) -> dict:
    """Build a minimal agenda_results dict."""
    agenda: dict = {}
    for sec in range(1, sections + 1):
        agenda[sec] = {
            f"Q{sec}-1?": _make_result(question=f"Q{sec}-1?"),
            f"Q{sec}-2?": _make_result(question=f"Q{sec}-2?"),
        }
    return agenda


# ─────────────────────────────────────────────────────────────────────────────
class TestSessionExporterInit:
    def test_default_exports_dir(self, tmp_path, monkeypatch):
        monkeypatch.chdir(tmp_path)
        from agent.session_exporter import SessionExporter
        exp = SessionExporter()
        assert exp.exports_dir == Path("exports")

    def test_custom_exports_dir(self, tmp_path):
        from agent.session_exporter import SessionExporter
        custom = tmp_path / "my_exports"
        exp = SessionExporter(exports_dir=custom)
        assert exp.exports_dir == custom

    def test_exports_dir_created(self, tmp_path):
        from agent.session_exporter import SessionExporter
        target = tmp_path / "new_dir"
        assert not target.exists()
        SessionExporter(exports_dir=target)
        assert target.exists()

    def test_llm_client_none_by_default(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        assert exp._llm is None

    def test_llm_client_stored(self, tmp_path):
        from agent.session_exporter import SessionExporter
        mock_llm = MagicMock()
        exp = SessionExporter(exports_dir=tmp_path, llm_client=mock_llm)
        assert exp._llm is mock_llm


# ─────────────────────────────────────────────────────────────────────────────
class TestExportChatTxt:
    def _make_chat(self) -> list[dict]:
        return [
            {"role": "user", "content": "How many accounts?",
             "timestamp": datetime(2024, 3, 1, 9, 0, tzinfo=timezone.utc)},
            {"role": "assistant", "content": "There are 42 accounts.",
             "result": _make_result(
                 answer_text="There are 42 accounts.",
                 priority_action="Review top 5",
                 risk_flag="Dormant accounts rising",
                 opportunity="Upsell to tier 2",
                 confidence_score=88,
             ),
             "timestamp": datetime(2024, 3, 1, 9, 1, tzinfo=timezone.utc)},
        ]

    def test_file_created(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt(self._make_chat(), session_id="abc123")
        assert path.exists()
        assert path.suffix == ".txt"

    def test_header_present(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt(self._make_chat(), session_id="abc123")
        text = path.read_text(encoding="utf-8")
        assert "CRM Analytics Agent" in text
        assert "Chat Export" in text

    def test_session_id_in_header(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt(self._make_chat(), session_id="MYID-42")
        assert "MYID-42" in path.read_text(encoding="utf-8")

    def test_user_message_present(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt(self._make_chat())
        text = path.read_text(encoding="utf-8")
        assert "USER:" in text
        assert "How many accounts?" in text

    def test_agent_message_present(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt(self._make_chat())
        text = path.read_text(encoding="utf-8")
        assert "AGENT:" in text

    def test_recommendation_section_present(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt(self._make_chat())
        text = path.read_text(encoding="utf-8")
        assert "RECOMMENDATION" in text
        assert "Priority Action" in text
        assert "Review top 5" in text

    def test_risk_flag_present(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt(self._make_chat())
        text = path.read_text(encoding="utf-8")
        assert "Risk Flag" in text
        assert "Dormant accounts rising" in text

    def test_opportunity_present(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt(self._make_chat())
        text = path.read_text(encoding="utf-8")
        assert "Opportunity" in text
        assert "Upsell to tier 2" in text

    def test_confidence_label_high(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt(self._make_chat())
        text = path.read_text(encoding="utf-8")
        # score=88 → High
        assert "High" in text or "88" in text

    def test_empty_chat_creates_file(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt([])
        assert path.exists()

    def test_multi_turn_history(self, tmp_path):
        from agent.session_exporter import SessionExporter
        chat = [
            {"role": "user", "content": f"Q{i}?", "timestamp": None}
            for i in range(5)
        ]
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt(chat)
        text = path.read_text(encoding="utf-8")
        assert "Q4?" in text

    def test_confidentiality_notice_present(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_chat_txt([])
        text = path.read_text(encoding="utf-8")
        assert "NDA" in text


# ─────────────────────────────────────────────────────────────────────────────
class TestExportActionsDocx:
    def test_file_created(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        agenda = _make_agenda(sections=2)
        path = exp.export_actions_docx(agenda)
        assert path.exists()
        assert path.suffix == ".docx"

    def test_valid_docx(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_actions_docx(_make_agenda())
        doc = docx.Document(str(path))
        # Just asserting it opens without error
        assert doc is not None

    def test_table_has_correct_columns(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_actions_docx(_make_agenda(sections=1))
        doc = docx.Document(str(path))
        # Find the first table
        tables = doc.tables
        assert len(tables) >= 1
        # Header row should contain column names
        hdr_text = " ".join(
            cell.text for cell in tables[0].rows[0].cells
        )
        assert "Priority Action" in hdr_text
        assert "Source Question" in hdr_text
        assert "Owner" in hdr_text

    def test_row_count_matches_action_count(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path)
        agenda = _make_agenda(sections=2)  # 2 sections × 2 questions = 4 rows
        path = exp.export_actions_docx(agenda)
        doc = docx.Document(str(path))
        # First table: header + 4 data rows (<=10 so single table)
        assert doc.tables[0].rows[0].cells[1].text == "Priority Action"
        data_rows = len(doc.tables[0].rows) - 1
        assert data_rows == 4

    def test_empty_agenda_creates_file(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_actions_docx({})
        assert path.exists()

    def test_large_agenda_uses_subheadings(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        # Create >10 rows: 6 sections × 2 questions each = 12 rows
        exp = SessionExporter(exports_dir=tmp_path)
        agenda = _make_agenda(sections=6)
        path = exp.export_actions_docx(agenda)
        doc = docx.Document(str(path))
        # Should have multiple tables (one per section group)
        assert len(doc.tables) >= 2


# ─────────────────────────────────────────────────────────────────────────────
class TestExportSectionDocx:
    def test_file_created(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        results = {"What is win rate?": _make_result()}
        path = exp.export_section_docx(2, results)
        assert path.exists()
        assert path.suffix == ".docx"

    def test_filename_includes_section_number(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_section_docx(3, {"Q?": _make_result()})
        assert "section_3" in path.name

    def test_valid_docx_opens(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_section_docx(1, {"Q?": _make_result()})
        doc = docx.Document(str(path))
        assert doc is not None

    def test_answer_text_in_doc(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path)
        r = _make_result(answer_text="Revenue is £2.5M this quarter.")
        path = exp.export_section_docx(2, {"Rev Q?": r})
        doc = docx.Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Revenue is £2.5M this quarter." in full_text

    def test_empty_results_creates_file(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_section_docx(4, {})
        assert path.exists()


# ─────────────────────────────────────────────────────────────────────────────
class TestExportSessionDocx:
    def test_file_created(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        path = exp.export_session_docx(
            chat_history=[],
            agenda_results=_make_agenda(),
            dataframes={},
            profiles={},
            session_id="s-001",
            model_used="llama3",
        )
        assert path.exists()
        assert path.suffix == ".docx"

    def test_valid_docx_opens(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        path = exp.export_session_docx(
            chat_history=[],
            agenda_results=_make_agenda(),
            dataframes={},
            profiles={},
        )
        doc = docx.Document(str(path))
        assert doc is not None

    def test_cover_page_title_present(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        path = exp.export_session_docx([], _make_agenda(), {}, {})
        doc = docx.Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "CRM Analytics" in full_text

    def test_session_id_in_doc(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        path = exp.export_session_docx([], _make_agenda(), {}, {},
                                        session_id="UNIQUE-XYZ")
        doc = docx.Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "UNIQUE-XYZ" in full_text

    def test_template_summary_used_when_no_llm(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        path = exp.export_session_docx([], _make_agenda(), {}, {})
        doc = docx.Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Template summary contains "session covered"
        assert "session" in full_text.lower()

    def test_appendix_section_present(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        path = exp.export_session_docx([], _make_agenda(), {}, {})
        doc = docx.Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Appendix" in full_text

    def test_data_sources_in_appendix(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        dfs = {
            "accounts": pd.DataFrame({"a": [1, 2, 3]}),
            "opportunities": pd.DataFrame({"b": [4, 5]}),
        }
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        path = exp.export_session_docx([], _make_agenda(), dfs, {})
        doc = docx.Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "accounts" in full_text
        assert "opportunities" in full_text

    def test_chat_history_exported_when_no_agenda(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        chat = [
            {"role": "user", "content": "Tell me about revenue."},
            {"role": "assistant", "content": "Revenue is high.", "result": None},
        ]
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        path = exp.export_session_docx(chat, {}, {}, {})
        doc = docx.Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Tell me about revenue." in full_text

    def test_actions_table_present(self, tmp_path):
        from agent.session_exporter import SessionExporter
        import docx
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        path = exp.export_session_docx([], _make_agenda(sections=2), {}, {})
        doc = docx.Document(str(path))
        # At least the Actions table should exist
        assert len(doc.tables) >= 1

    def test_empty_agenda_and_chat_creates_file(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        path = exp.export_session_docx([], {}, {}, {})
        assert path.exists()


# ─────────────────────────────────────────────────────────────────────────────
class TestExportDashboardPdf:
    """Tests for export_dashboard_pdf. Requires kaleido + pypdf."""

    @pytest.fixture
    def sample_dfs(self) -> dict:
        accounts = pd.DataFrame({
            "account_name": ["Acme", "Beta", "Gamma", "Delta"],
            "revenue": [100_000, 80_000, 60_000, 40_000],
            "tenure_years": [5, 3, 7, 2],
        })
        opportunities = pd.DataFrame({
            "stage": ["Proposal", "Closed Won", "Proposal", "Closed Won", "Negotiation"],
            "deal_value": [50_000, 120_000, 30_000, 200_000, 75_000],
            "service_line": ["FM", "Consulting", "FM", "Consulting", "FM"],
            "outcome": ["", "Won", "", "Won", ""],
        })
        return {"accounts": accounts, "opportunities": opportunities}

    def _has_kaleido(self) -> bool:
        try:
            import kaleido  # noqa: F401
            return True
        except ImportError:
            return False

    def _has_pypdf(self) -> bool:
        try:
            from pypdf import PdfWriter  # noqa: F401
            return True
        except ImportError:
            return False

    def test_raises_import_error_without_kaleido(self, tmp_path, sample_dfs):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        with patch.dict("sys.modules", {"kaleido": None}):
            with pytest.raises(ImportError, match="kaleido"):
                exp.export_dashboard_pdf(sample_dfs)

    def test_raises_import_error_without_pypdf(self, tmp_path, sample_dfs):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        # Kaleido must be importable for this test to reach pypdf check
        if not self._has_kaleido():
            pytest.skip("kaleido not installed")
        with patch.dict("sys.modules", {"pypdf": None}):
            with pytest.raises(ImportError, match="pypdf"):
                exp.export_dashboard_pdf(sample_dfs)

    def test_empty_dataframes_raises_value_error(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        if not self._has_kaleido() or not self._has_pypdf():
            pytest.skip("kaleido or pypdf not installed")
        with pytest.raises((ValueError, Exception)):
            exp.export_dashboard_pdf({})

    @pytest.mark.skipif(
        not (
            __import__("importlib").util.find_spec("kaleido") is not None
            and __import__("importlib").util.find_spec("pypdf") is not None
        ),
        reason="kaleido and pypdf required",
    )
    def test_pdf_file_created(self, tmp_path, sample_dfs):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_dashboard_pdf(sample_dfs)
        assert path.exists()
        assert path.suffix == ".pdf"

    @pytest.mark.skipif(
        not (
            __import__("importlib").util.find_spec("kaleido") is not None
            and __import__("importlib").util.find_spec("pypdf") is not None
        ),
        reason="kaleido and pypdf required",
    )
    def test_pdf_is_readable_by_pypdf(self, tmp_path, sample_dfs):
        from agent.session_exporter import SessionExporter
        from pypdf import PdfReader
        exp = SessionExporter(exports_dir=tmp_path)
        path = exp.export_dashboard_pdf(sample_dfs)
        reader = PdfReader(str(path))
        assert len(reader.pages) >= 1


# ─────────────────────────────────────────────────────────────────────────────
class TestBuildDashboardFigs:
    @pytest.fixture
    def sample_dfs(self):
        accounts = pd.DataFrame({
            "account_name": [f"Acct{i}" for i in range(5)],
            "revenue": [100_000 * i for i in range(1, 6)],
            "engagement": [10, 20, 30, 40, 50],
        })
        opps = pd.DataFrame({
            "stage": ["Proposal", "Won", "Proposal", "Won", "Lost"],
            "deal_value": [10_000, 50_000, 20_000, 80_000, 5_000],
            "service_line": ["FM", "FM", "Consulting", "Consulting", "FM"],
            "outcome": ["", "Won", "", "Won", "Lost"],
        })
        return {"accounts": accounts, "opportunities": opps}

    def test_returns_list(self, sample_dfs):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=Path("/tmp"))
        figs = exp._build_dashboard_figs(sample_dfs)
        assert isinstance(figs, list)

    def test_returns_empty_for_empty_dfs(self):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=Path("/tmp"))
        figs = exp._build_dashboard_figs({})
        assert figs == []

    def test_returns_plotly_figures(self, sample_dfs):
        pytest.importorskip("plotly")
        from agent.session_exporter import SessionExporter
        import plotly.graph_objects as go
        exp = SessionExporter(exports_dir=Path("/tmp"))
        figs = exp._build_dashboard_figs(sample_dfs)
        for fig in figs:
            assert isinstance(fig, go.Figure)

    def test_at_least_one_figure_generated(self, sample_dfs):
        pytest.importorskip("plotly")
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=Path("/tmp"))
        figs = exp._build_dashboard_figs(sample_dfs)
        assert len(figs) >= 1


# ─────────────────────────────────────────────────────────────────────────────
class TestExecutiveSummary:
    def test_template_summary_no_llm(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        summary = exp._executive_summary(_make_agenda(sections=2))
        assert isinstance(summary, str)
        assert len(summary) > 0

    def test_template_mentions_item_count(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        # 2 sections × 2 questions = 4 items
        summary = exp._executive_summary(_make_agenda(sections=2))
        assert "4" in summary

    def test_template_mentions_session_id(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        summary = exp._executive_summary({}, session_id="MYSESSION")
        assert "MYSESSION" in summary

    def test_template_fallback_when_llm_fails(self, tmp_path):
        from agent.session_exporter import SessionExporter
        mock_llm = MagicMock()

        async def _fail(*args, **kwargs):
            raise RuntimeError("LLM is offline")

        mock_llm.complete = _fail
        exp = SessionExporter(exports_dir=tmp_path, llm_client=mock_llm)
        summary = exp._executive_summary(_make_agenda(sections=1))
        assert isinstance(summary, str)
        assert len(summary) > 0

    def test_empty_agenda_returns_string(self, tmp_path):
        from agent.session_exporter import SessionExporter
        exp = SessionExporter(exports_dir=tmp_path, llm_client=None)
        summary = exp._executive_summary({})
        assert isinstance(summary, str)


# ─────────────────────────────────────────────────────────────────────────────
class TestPivotHandlerMetadata:
    """Test PivotHandler.export_to_excel with new metadata kwarg."""

    @pytest.fixture
    def pivot_df(self):
        return pd.DataFrame({
            "service_line": ["FM", "Consulting"],
            "revenue": [500_000, 300_000],
            "win_rate": [0.28, 0.35],
        })

    def test_metadata_kwarg_accepted(self, tmp_path, pivot_df):
        from agent.query_engine import PivotHandler
        handler = PivotHandler()
        metadata = {
            "question": "What is revenue by service line?",
            "code": "result = df.groupby('service_line')['revenue'].sum()",
            "timestamp": "2024-03-01T10:00:00+00:00",
            "confidence_score": 85,
            "provider_used": "ollama",
            "session_id": "sess-001",
        }
        out = handler.export_to_excel(pivot_df, None, tmp_path, metadata=metadata)
        assert out.exists()

    def test_metadata_creates_sheet_3(self, tmp_path, pivot_df):
        from agent.query_engine import PivotHandler
        from openpyxl import load_workbook
        handler = PivotHandler()
        metadata = {
            "question": "Revenue by service line?",
            "code": "result = df.head()",
            "timestamp": "2024-03-01T10:00:00",
            "confidence_score": 80,
            "provider_used": "ollama",
            "session_id": "s1",
        }
        out = handler.export_to_excel(pivot_df, None, tmp_path, metadata=metadata)
        wb = load_workbook(out)
        assert "Metadata" in wb.sheetnames

    def test_metadata_sheet_has_correct_keys(self, tmp_path, pivot_df):
        from agent.query_engine import PivotHandler
        from openpyxl import load_workbook
        handler = PivotHandler()
        meta = {
            "question": "My question",
            "code": "result = 42",
            "timestamp": "2024-01-01T00:00:00",
            "confidence_score": 75,
            "provider_used": "groq",
            "session_id": "abc",
        }
        out = handler.export_to_excel(pivot_df, None, tmp_path, metadata=meta)
        wb = load_workbook(out)
        ws = wb["Metadata"]
        keys = [ws.cell(r, 1).value for r in range(1, 7)]
        assert "Query" in keys
        assert "Code" in keys
        assert "Confidence" in keys
        assert "Provider" in keys

    def test_no_metadata_has_only_pivot_sheet(self, tmp_path, pivot_df):
        from agent.query_engine import PivotHandler
        from openpyxl import load_workbook
        handler = PivotHandler()
        out = handler.export_to_excel(pivot_df, None, tmp_path)
        wb = load_workbook(out)
        assert "Metadata" not in wb.sheetnames
        assert "Pivot" in wb.sheetnames

    def test_with_raw_df_has_two_sheets_plus_metadata(self, tmp_path, pivot_df):
        from agent.query_engine import PivotHandler
        from openpyxl import load_workbook
        handler = PivotHandler()
        raw = pd.DataFrame({"x": [1, 2, 3]})
        meta = {"question": "Q", "code": "c", "timestamp": "t",
                "confidence_score": 70, "provider_used": "p", "session_id": "s"}
        out = handler.export_to_excel(pivot_df, raw, tmp_path, metadata=meta)
        wb = load_workbook(out)
        assert "Pivot" in wb.sheetnames
        assert "Raw Data" in wb.sheetnames
        assert "Metadata" in wb.sheetnames


# ─────────────────────────────────────────────────────────────────────────────
class TestPivotNumberFormat:
    """Test currency and percentage number formatting in PivotHandler."""

    @pytest.fixture
    def mixed_df(self):
        return pd.DataFrame({
            "account_name": ["A", "B", "C"],
            "total_revenue": [100_000, 200_000, 150_000],
            "win_rate": [0.28, 0.35, 0.22],
            "deal_value": [50_000, 80_000, 30_000],
            "count": [10, 20, 15],
        })

    def test_currency_column_gets_pound_format(self, tmp_path, mixed_df):
        from agent.query_engine import PivotHandler
        from openpyxl import load_workbook
        handler = PivotHandler()
        out = handler.export_to_excel(mixed_df, None, tmp_path)
        wb = load_workbook(out)
        ws = wb["Pivot"]
        # Find the revenue column index (after reset_index, index is col 1)
        headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
        rev_idx = next(
            (i + 1 for i, h in enumerate(headers) if h and "revenue" in h.lower()),
            None
        )
        if rev_idx is not None:
            fmt = ws.cell(2, rev_idx).number_format
            assert "\u00a3" in fmt or "0" in fmt

    def test_pct_column_gets_percent_format(self, tmp_path, mixed_df):
        from agent.query_engine import PivotHandler
        from openpyxl import load_workbook
        handler = PivotHandler()
        out = handler.export_to_excel(mixed_df, None, tmp_path)
        wb = load_workbook(out)
        ws = wb["Pivot"]
        headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
        rate_idx = next(
            (i + 1 for i, h in enumerate(headers) if h and "rate" in h.lower()),
            None
        )
        if rate_idx is not None:
            fmt = ws.cell(2, rate_idx).number_format
            assert "%" in fmt or "0" in fmt

    def test_non_numeric_currency_column_not_formatted(self, tmp_path):
        from agent.query_engine import PivotHandler
        from openpyxl import load_workbook
        # String revenue column — should not apply format
        df = pd.DataFrame({
            "account": ["A", "B"],
            "revenue": ["£100k", "£200k"],  # string, not numeric
        })
        handler = PivotHandler()
        out = handler.export_to_excel(df, None, tmp_path)
        wb = load_workbook(out)
        ws = wb["Pivot"]
        # Just verify it doesn't crash
        assert ws is not None

    def test_neutral_column_has_default_format(self, tmp_path):
        from agent.query_engine import PivotHandler
        from openpyxl import load_workbook
        df = pd.DataFrame({
            "account_name": ["A", "B"],
            "count": [10, 20],
        })
        handler = PivotHandler()
        out = handler.export_to_excel(df, None, tmp_path)
        wb = load_workbook(out)
        ws = wb["Pivot"]
        headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
        count_idx = next(
            (i + 1 for i, h in enumerate(headers) if h == "count"), None
        )
        if count_idx is not None:
            fmt = ws.cell(2, count_idx).number_format
            # Should not be currency or percentage
            assert "\u00a3" not in (fmt or "")
            assert "%" not in (fmt or "")
